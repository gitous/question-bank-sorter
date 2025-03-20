import os
import re
import time
import pandas as pd
import pypinyin
import docx
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor, Pt
import pdfplumber
from openpyxl import Workbook
from concurrent.futures import ThreadPoolExecutor
import tkinter as tk
from tkinter import filedialog, messagebox, StringVar, Label, scrolledtext
import threading
import queue

class Logger:
    def __init__(self, text_widget):
        self.text_widget = text_widget
        self.queue = queue.Queue()
        self.update_interval = 100  # ms
        
    def log(self, message):
        """添加日志消息到队列"""
        self.queue.put(message)
        
    def update_display(self):
        """从队列获取消息并更新UI"""
        try:
            while True:
                message = self.queue.get_nowait()
                self.text_widget.insert(tk.END, message + "\n")
                self.text_widget.see(tk.END)
                self.queue.task_done()
        except queue.Empty:
            self.text_widget.after(self.update_interval, self.update_display)
            
    def start(self):
        """开始日志更新循环"""
        self.text_widget.after(self.update_interval, self.update_display)

class QuestionBankSorter:
    def __init__(self, logger=None):
        self.questions = []
        self.logger = logger
        
        # 改进的题目和答案识别模式
        self.question_pattern = re.compile(r'(?:(?:\d+[\.\s、]+)|^)(.+?)(?:(?=\d+[\.\s、]+)|$)', re.DOTALL)
        self.answer_pattern = re.compile(r'(?:标准答案[:：]?\s*|答案[:：]?\s*|[\(（]答案[\)）][:：]?\s*)([a-dA-D]|正确|错误)', re.IGNORECASE)
        
    def log(self, message):
        """输出日志消息"""
        if self.logger:
            self.logger.log(message)
        else:
            print(message)
        
    def get_first_char_pinyin(self, text):
        """获取文本中第一个汉字的拼音"""
        for char in text:
            if '\u4e00' <= char <= '\u9fff':  # 检查字符是否为汉字
                pinyin_list = pypinyin.lazy_pinyin(char, style=pypinyin.NORMAL)
                if pinyin_list:
                    return pinyin_list[0].lower()  # 转为小写以确保排序正确
        return "z"  # 非汉字起始的文本默认排在最后
    
    def clean_question_text(self, text):
        """清理题目文本，去除题号"""
        # 去除题目开头的数字和标点
        cleaned_text = re.sub(r'^(\d+[\.\s、]+)', '', text.strip())
        return cleaned_text
    
    def extract_from_pdf(self, pdf_path):
        """从PDF文件中提取题目和答案"""
        self.log(f"正在处理PDF文件: {os.path.basename(pdf_path)}")
        temp_questions = []
        
        try:
            with pdfplumber.open(pdf_path) as pdf:
                text = ""
                total_pages = len(pdf.pages)
                
                for i, page in enumerate(pdf.pages):
                    self.log(f"正在处理PDF第 {i+1}/{total_pages} 页")
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    
            # 查找所有题目
            matches = self.question_pattern.findall(text)
            self.log(f"从PDF中找到 {len(matches)} 个潜在题目")
            
            for match in matches:
                question_text = match.strip()
                if not question_text:
                    continue
                    
                # 查找答案
                answer_match = self.answer_pattern.search(question_text)
                answer = answer_match.group(1).upper() if answer_match else ""
                
                # 清理题目文本
                cleaned_text = self.clean_question_text(question_text)
                
                if cleaned_text:
                    temp_questions.append({
                        'text': cleaned_text,
                        'answer': answer,
                        'sort_key': self.get_first_char_pinyin(cleaned_text)
                    })
                    
            self.log(f"成功从PDF提取 {len(temp_questions)} 个题目")
            return temp_questions
            
        except Exception as e:
            self.log(f"处理PDF文件时出错: {str(e)}")
            return []
    
    def extract_from_docx(self, docx_path):
        """从DOCX文件中提取题目和答案"""
        self.log(f"正在处理Word文件: {os.path.basename(docx_path)}")
        temp_questions = []
        
        try:
            doc = Document(docx_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            
            # 查找所有题目
            matches = self.question_pattern.findall(text)
            self.log(f"从Word中找到 {len(matches)} 个潜在题目")
            
            for match in matches:
                question_text = match.strip()
                if not question_text:
                    continue
                    
                # 查找答案
                answer_match = self.answer_pattern.search(question_text)
                answer = answer_match.group(1).upper() if answer_match else ""
                
                # 清理题目文本
                cleaned_text = self.clean_question_text(question_text)
                
                if cleaned_text:
                    temp_questions.append({
                        'text': cleaned_text,
                        'answer': answer,
                        'sort_key': self.get_first_char_pinyin(cleaned_text)
                    })
                    
            self.log(f"成功从Word提取 {len(temp_questions)} 个题目")
            return temp_questions
            
        except Exception as e:
            self.log(f"处理Word文件时出错: {str(e)}")
            return []
    
    def extract_from_txt(self, txt_path):
        """从TXT文件中提取题目和答案"""
        self.log(f"正在处理文本文件: {os.path.basename(txt_path)}")
        temp_questions = []
        
        try:
            with open(txt_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
            
            # 查找所有题目
            matches = self.question_pattern.findall(text)
            self.log(f"从TXT中找到 {len(matches)} 个潜在题目")
            
            for match in matches:
                question_text = match.strip()
                if not question_text:
                    continue
                    
                # 查找答案
                answer_match = self.answer_pattern.search(question_text)
                answer = answer_match.group(1).upper() if answer_match else ""
                
                # 清理题目文本
                cleaned_text = self.clean_question_text(question_text)
                
                if cleaned_text:
                    temp_questions.append({
                        'text': cleaned_text,
                        'answer': answer,
                        'sort_key': self.get_first_char_pinyin(cleaned_text)
                    })
                    
            self.log(f"成功从TXT提取 {len(temp_questions)} 个题目")
            return temp_questions
            
        except Exception as e:
            self.log(f"处理TXT文件时出错: {str(e)}")
            return []
    
    def extract_from_excel(self, excel_path):
        """从Excel文件中提取题目和答案"""
        self.log(f"正在处理Excel文件: {os.path.basename(excel_path)}")
        temp_questions = []
        
        try:
            df = pd.read_excel(excel_path)
            total_rows = len(df)
            self.log(f"Excel文件包含 {total_rows} 行数据")
            
            # 尝试识别题目和答案列
            possible_columns = {}
            for col in df.columns:
                col_str = str(col).lower()
                if '题' in col_str or '问题' in col_str or 'question' in col_str:
                    possible_columns['question'] = col
                elif '答案' in col_str or 'answer' in col_str:
                    possible_columns['answer'] = col
            
            # 处理每一行
            for i, row in df.iterrows():
                if i % 100 == 0:
                    self.log(f"正在处理Excel第 {i+1}/{total_rows} 行")
                
                # 如果找到了题目列，使用它
                if 'question' in possible_columns:
                    question_text = str(row[possible_columns['question']])
                # 否则尝试使用第一列作为题目
                else:
                    question_text = str(row[df.columns[0]])
                
                # 如果找到了答案列，使用它
                answer = ""
                if 'answer' in possible_columns:
                    answer = str(row[possible_columns['answer']]).upper()
                # 否则尝试从题目文本中提取答案
                else:
                    answer_match = self.answer_pattern.search(question_text)
                    if answer_match:
                        answer = answer_match.group(1).upper()
                
                # 清理题目文本
                cleaned_text = self.clean_question_text(question_text)
                
                if cleaned_text:
                    temp_questions.append({
                        'text': cleaned_text,
                        'answer': answer,
                        'sort_key': self.get_first_char_pinyin(cleaned_text)
                    })
            
            self.log(f"成功从Excel提取 {len(temp_questions)} 个题目")
            return temp_questions
            
        except Exception as e:
            self.log(f"处理Excel文件时出错: {str(e)}")
            return []
    
    def process_file(self, file_path):
        """根据文件扩展名处理文件"""
        _, ext = os.path.splitext(file_path.lower())
        
        if ext == '.pdf':
            return self.extract_from_pdf(file_path)
        elif ext == '.docx' or ext == '.doc':
            return self.extract_from_docx(file_path)
        elif ext == '.txt':
            return self.extract_from_txt(file_path)
        elif ext == '.xlsx' or ext == '.xls':
            return self.extract_from_excel(file_path)
        else:
            self.log(f"不支持的文件格式: {ext}")
            return []
    
    def process_files(self, file_paths):
        """并行处理多个文件"""
        self.questions = []
        total_files = len(file_paths)
        
        self.log(f"开始处理 {total_files} 个文件...")
        
        # 顺序处理文件以便更好地显示进度
        for i, file_path in enumerate(file_paths):
            self.log(f"处理文件 {i+1}/{total_files}: {os.path.basename(file_path)}")
            result = self.process_file(file_path)
            self.questions.extend(result)
            self.log(f"文件 {os.path.basename(file_path)} 处理完成，提取了 {len(result)} 个题目")
        
        # 按拼音排序题目
        self.log("正在按拼音排序题目...")
        self.questions.sort(key=lambda x: x['sort_key'])
        self.log(f"排序完成，共有 {len(self.questions)} 个题目")
    
    def save_to_word(self, output_path):
        """保存排序后的题目到Word文档"""
        self.log("正在创建Word文档...")
        try:
            doc = Document()
            
            # 设置文档标题
            title = doc.add_heading('题库 - 按拼音排序', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 设置默认字体为宋体
            style = doc.styles['Normal']
            font = style.font
            font.name = '宋体'
            font.size = Pt(12)
            
            # 添加每个题目
            current_letter = None
            for i, question in enumerate(self.questions, 1):
                # 如果是新的字母分组，添加分隔符
                first_letter = question['sort_key'][0] if question['sort_key'] else ''
                if first_letter != current_letter:
                    current_letter = first_letter
                    section = doc.add_heading(f'{current_letter.upper()} 部分', level=1)
                    section.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                
                # 添加题目
                q_para = doc.add_paragraph()
                q_para.add_run(f"{i}. {question['text']}").bold = True
                
                # 如果有答案，添加答案
                if question['answer']:
                    a_para = doc.add_paragraph()
                    answer_run = a_para.add_run(f"答案: {question['answer']}")
                    answer_run.font.color.rgb = RGBColor(0, 0, 128)  # 深蓝色
                
                # 添加空行
                doc.add_paragraph()
                
                if i % 50 == 0:
                    self.log(f"已处理 {i}/{len(self.questions)} 个题目")
            
            # 保存文档
            doc.save(output_path)
            self.log(f"Word文档已保存到 {output_path}")
            return True
        except Exception as e:
            self.log(f"保存Word文档时出错: {str(e)}")
            return False
    
    def save_to_excel(self, output_path):
        """保存排序后的题目到Excel表格"""
        self.log("正在创建Excel文件...")
        try:
            # 创建数据帧
            data = {
                '题号': [i for i in range(1, len(self.questions) + 1)],
                '拼音首字母': [q['sort_key'][0].upper() if q['sort_key'] else '' for q in self.questions],
                '题目': [q['text'] for q in self.questions],
                '答案': [q['answer'] for q in self.questions]
            }
            
            df = pd.DataFrame(data)
            
            # 保存到Excel
            df.to_excel(output_path, index=False, sheet_name='题库')
            self.log(f"Excel文件已保存到 {output_path}")
            return True
        except Exception as e:
            self.log(f"保存Excel文件时出错: {str(e)}")
            return False
    
    def save_to_txt(self, output_path):
        """保存排序后的题目到文本文件"""
        self.log("正在创建文本文件...")
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                current_letter = None
                
                for i, question in enumerate(self.questions, 1):
                    # 如果是新的字母分组，添加分隔符
                    first_letter = question['sort_key'][0] if question['sort_key'] else ''
                    if first_letter != current_letter:
                        current_letter = first_letter
                        f.write(f"\n{'='*50}\n{current_letter.upper()} 部分\n{'='*50}\n\n")
                    
                    # 写入题目
                    f.write(f"{i}. {question['text']}\n")
                    
                    # 如果有答案，添加答案
                    if question['answer']:
                        f.write(f"答案: {question['answer']}\n")
                    
                    # 添加空行
                    f.write("\n")
                    
                    if i % 100 == 0:
                        self.log(f"已写入 {i}/{len(self.questions)} 个题目")
            
            self.log(f"文本文件已保存到 {output_path}")
            return True
        except Exception as e:
            self.log(f"保存文本文件时出错: {str(e)}")
            return False
        
    def save_all_formats(self, base_output_path):
        """保存为所有格式(Word, Excel, TXT)"""
        results = {
            "Word": self.save_to_word(f"{base_output_path}.docx"),
            "Excel": self.save_to_excel(f"{base_output_path}.xlsx"),
            "Text": self.save_to_txt(f"{base_output_path}.txt")
        }
        
        success_count = sum(1 for result in results.values() if result)
        if success_count == 3:
            self.log("全部文件保存成功！")
        else:
            self.log(f"部分文件保存失败: 成功率 {success_count}/3")
        
        return results

def run_processor(file_paths, output_path, logger):
    """在新线程中运行处理函数"""
    try:
        sorter = QuestionBankSorter(logger=logger)
        sorter.process_files(file_paths)
        results = sorter.save_all_formats(output_path)
        
        # 检查是否有成功的保存
        if any(results.values()):
            success_formats = [fmt for fmt, result in results.items() if result]
            logger.log(f"题库已成功排序并保存为以下格式：{', '.join(success_formats)}")
        else:
            logger.log("所有保存操作均失败，请检查文件路径和权限")
    
    except Exception as e:
        logger.log(f"处理过程中出现错误: {str(e)}")

def main():
    """主函数，创建GUI界面"""
    root = tk.Tk()
    root.title("题库排序工具")
    root.geometry("700x500")
    
    file_paths = []
    status_var = StringVar(value="准备就绪，请添加文件")
    
    def add_files():
        paths = filedialog.askopenfilenames(
            title="选择题库文件",
            filetypes=[
                ("所有支持的文件", "*.pdf;*.docx;*.doc;*.txt;*.xlsx;*.xls"),
                ("PDF文件", "*.pdf"),
                ("Word文件", "*.docx;*.doc"),
                ("文本文件", "*.txt"),
                ("Excel文件", "*.xlsx;*.xls")
            ]
        )
        if paths:
            for path in paths:
                if path not in file_paths:
                    file_paths.append(path)
                    files_listbox.insert(tk.END, os.path.basename(path))
            status_var.set(f"已添加 {len(file_paths)} 个文件")
    
    def remove_file():
        selection = files_listbox.curselection()
        if selection:
            index = selection[0]
            files_listbox.delete(index)
            file_paths.pop(index)
            status_var.set(f"已添加 {len(file_paths)} 个文件")
    
    def process_files_callback():
        if not file_paths:
            messagebox.showwarning("警告", "请先添加文件")
            return
            
        output_path = filedialog.asksaveasfilename(
            title="保存排序后的题库",
            filetypes=[("所有文件", "*")]
        )
        
        if output_path:
            # 禁用按钮，防止重复点击
            add_button.config(state=tk.DISABLED)
            remove_button.config(state=tk.DISABLED)
            process_button.config(state=tk.DISABLED)
            
            # 清空日志
            log_text.delete(1.0, tk.END)
            
            # 启动处理线程
            process_thread = threading.Thread(
                target=run_processor,
                args=(file_paths, output_path, logger),
                daemon=True
            )
            process_thread.start()
            
            # 开始定期检查线程是否完成
            check_thread_completion(process_thread)
    
    def check_thread_completion(thread):
        if thread.is_alive():
            # 如果线程仍在运行，100ms后再次检查
            root.after(100, lambda: check_thread_completion(thread))
        else:
            # 线程完成，重新启用按钮
            add_button.config(state=tk.NORMAL)
            remove_button.config(state=tk.NORMAL)
            process_button.config(state=tk.NORMAL)
            status_var.set("处理完成！")
    
    # 创建GUI元素
    # 顶部框架
    frame_top = tk.Frame(root)
    frame_top.pack(fill=tk.X, padx=10, pady=10)
    
    add_button = tk.Button(frame_top, text="添加文件", command=add_files, width=15)
    add_button.pack(side=tk.LEFT, padx=5)
    
    remove_button = tk.Button(frame_top, text="移除文件", command=remove_file, width=15)
    remove_button.pack(side=tk.LEFT, padx=5)
    
    # 文件列表框架
    frame_files = tk.Frame(root)
    frame_files.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
    
    files_label = tk.Label(frame_files, text="已添加的文件:")
    files_label.pack(anchor=tk.W)
    
    files_listbox = tk.Listbox(frame_files, height=6)
    files_listbox.pack(fill=tk.X, expand=False)
    
    # 日志框架
    frame_log = tk.Frame(root)
    frame_log.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
    
    log_label = tk.Label(frame_log, text="处理日志:")
    log_label.pack(anchor=tk.W)
    
    log_text = scrolledtext.ScrolledText(frame_log, height=12)
    log_text.pack(fill=tk.BOTH, expand=True)
    
    # 创建日志记录器
    logger = Logger(log_text)
    logger.start()
    
    # 状态栏
    status_bar = Label(root, textvariable=status_var, bd=1, relief=tk.SUNKEN, anchor=tk.W)
    status_bar.pack(side=tk.BOTTOM, fill=tk.X)
    
    # 处理按钮
    process_button = tk.Button(root, text="处理并保存", command=process_files_callback, width=20, height=2)
    process_button.pack(pady=10)
    
    root.mainloop()

if __name__ == "__main__":
    main()
